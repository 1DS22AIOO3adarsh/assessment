from typing import List
from langchain.docstore.document import Document
from io import BytesIO
import pdfplumber
import docx

def load_document(file) -> List[Document]:
    """Load a single file and return list of Documents"""
    docs = []
    file_type = file.type

    if "pdf" in file_type:
        with pdfplumber.open(BytesIO(file.read())) as pdf:
            text = "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())
        docs.append(Document(page_content=text))

    elif "word" in file_type or file.name.endswith(".docx"):
        doc = docx.Document(BytesIO(file.read()))
        text = "\n".join([para.text for para in doc.paragraphs])
        docs.append(Document(page_content=text))

    elif "text" in file_type or file.name.endswith(".txt"):
        text = file.read().decode("utf-8")
        docs.append(Document(page_content=text))

    else:
        # unsupported filetype
        pass

    return docs
