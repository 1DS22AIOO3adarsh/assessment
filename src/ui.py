import streamlit as st
from .loaders import load_document
from .processor import split_documents
from .gemini_chain import RAGChat
from docx import Document
from io import BytesIO


def render_app():
    st.set_page_config(page_title="RAG Streamlit App", layout="wide")
    st.title("Chat with Your Documents")

    # File uploader
    files = st.file_uploader(
        "Upload PDF, DOCX, or TXT",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True
    )

    # Initialize state once
    if 'chat_history' not in st.session_state:
        st.session_state.chat_history = []
    if 'rag_chat' not in st.session_state:
        st.session_state.rag_chat = None
    if 'doc_keys' not in st.session_state:
        st.session_state.doc_keys = []

    # If new files are uploaded or app loaded for first time
    if files:
        current_keys = [file.name + str(file.size) for file in files]
        if current_keys != st.session_state.doc_keys:
            st.session_state.doc_keys = current_keys

            # Load and process files
            docs = []
            for file in files:
                docs.extend(load_document(file))

            with st.spinner("Processing documents..."):
                chunks = split_documents(docs)
                st.session_state.rag_chat = RAGChat(chunks)
                st.success("Documents processed successfully!")

    # Chat input and interaction
    query = st.chat_input("Ask a question:")
    if query:
        st.session_state.chat_history.append(("user", query))

        if st.session_state.rag_chat:
            answer = st.session_state.rag_chat(query)
        else:
            answer = "Document processing is not completed yet."

        st.session_state.chat_history.append(("assistant", answer))

    # Show chat history
    for sender, message in st.session_state.chat_history:
        with st.chat_message(sender):
            st.markdown(message)


    def save_chat_to_docx(chat_history):
        doc = Document()
        doc.add_heading("Chat Transcript", level=1)

        for sender, message in chat_history:
            if sender == "user":
                doc.add_paragraph("User:", style="Heading 2")
            else:
                doc.add_paragraph("Assistant:", style="Heading 2")
            doc.add_paragraph(message, style="Normal")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

        # Show download button if chat exists
    if st.session_state.chat_history:
        docx_buffer = save_chat_to_docx(st.session_state.chat_history)
        st.download_button(
            label="Download Chat as .docx",
            data=docx_buffer,
            file_name="chat_transcript.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
